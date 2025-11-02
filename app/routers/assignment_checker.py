# app/routers/assignment_checker.py
from pathlib import Path
from fastapi import APIRouter, Request, UploadFile, File, Form
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.templating import Jinja2Templates

from app.services.assignment_checker_agents import (
    IngestionAgent,
    RequirementAgent,
    ScoringAgent,
    FeedbackAgent,
)
from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from dotenv import load_dotenv
import docx, os, numpy as np
from openai import OpenAI

router = APIRouter()

PROJECT_ROOT = Path(__file__).resolve().parents[2]
templates = Jinja2Templates(directory=str(PROJECT_ROOT / "templates"))
OPENAI_API_KEY='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'

UPLOAD_DIR = PROJECT_ROOT / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

client = OpenAI(api_key=OPENAI_API_KEY)

ingestion_agent = IngestionAgent(upload_dir=UPLOAD_DIR)
requirement_agent = RequirementAgent()
scoring_agent = ScoringAgent()
feedback_agent = FeedbackAgent()

def extract_text_and_font(path):
    """Extract text (and detect font if DOCX)."""
    ext = os.path.splitext(path)[1].lower()
    text, font = "", "Unknown"

    try:
        if ext == ".pdf":
            reader = PdfReader(path)
            text = "\n".join([p.extract_text() or "" for p in reader.pages])
        elif ext == ".docx":
            doc = docx.Document(path)
            text = "\n".join([p.text for p in doc.paragraphs])
            # detect font if available
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.font.name:
                        font = run.font.name
                        break
                if font != "Unknown":
                    break
        elif ext == ".txt":
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()
    except Exception as e:
        print("[ERROR extracting]", e)
    return text.strip(), font


def get_embedding(text):
    """Compute OpenAI embedding."""
    try:
        resp = client.embeddings.create(model="text-embedding-3-small", input=text[:2000])
        return np.array(resp.data[0].embedding)
    except Exception as e:
        print("[ERROR] Embedding:", e)
        return np.zeros((1536,))


def compare_texts(req_text, sub_text):
    """Compute cosine similarity."""
    a, b = get_embedding(req_text), get_embedding(sub_text)
    if np.linalg.norm(a) == 0 or np.linalg.norm(b) == 0:
        return 0.0
    sim = np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))
    return round(float(sim) * 100, 2)


def generate_feedback(req_text, sub_text):
    """Use GPT to generate feedback."""
    prompt = f"""
    Professor Requirements:
    {req_text}

    Student Submission (first 1500 chars):
    {sub_text[:1500]}

    Evaluate how well the submission meets the requirements.
    Mention positives, negatives, and give clear improvement suggestions.
    """
    try:
        res = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a grading assistant that highlights issues."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.5,
            max_tokens=250
        )
        return res.choices[0].message.content.strip()
    except Exception as e:
        print("[ERROR Feedback]", e)
        return "Unable to generate feedback (OpenAI error)."


@router.get("/assignment-checker", response_class=HTMLResponse)
async def assignment_checker_page(request: Request):
    return templates.TemplateResponse(
        "assignment_checker.html",
        {"request": request},
    )


@router.post("/api/assignment-checker/score")
async def assignment_checker_api(
    request: Request,
    requirements: str = Form(...),
    student_text: str = Form(""),
    file: UploadFile | None = File(None),
):
    eq_text = request.form.get("detailsText", "").strip()
    req_file = request.files.get("detailsFile")
    sub_file = request.files.get("submission")

    if not sub_file:
        return jsonify({"error": "Please upload the student's submission file."}), 400

    # save + extract
    if req_file:
        req_path = os.path.join(UPLOAD_DIR, secure_filename(req_file.filename))
        req_file.save(req_path)
        req_text, _ = extract_text_and_font(req_path)

    sub_path = os.path.join(UPLOAD_DIR, secure_filename(sub_file.filename))
    sub_file.save(sub_path)
    sub_text, sub_font = extract_text_and_font(sub_path)

    if not req_text or not sub_text:
        return jsonify({"error": "Missing or unreadable text."}), 400

    # compute similarity
    score = compare_texts(req_text, sub_text)
    feedback = generate_feedback(req_text, sub_text)

    # detect issues
    issues = []
    if "times new roman" in req_text.lower() and "times new roman" not in sub_font.lower():
        issues.append(f"Font mismatch: expected Times New Roman, found {sub_font}.")
    if len(sub_text.split()) < 100:
        issues.append("Submission may be too short based on requirements.")

    # return full payload
    return jsonify({
        "similarity_score": score,
        "feedback": feedback,
        "issues": issues,
        "preview_text": sub_text[:1200] + ("..." if len(sub_text) > 1200 else "")
    })
